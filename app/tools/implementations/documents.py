from __future__ import annotations

import os
import tempfile
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from markdown import markdown
from pydantic import BaseModel, Field
from typing import Any

from app.core.storage import upload_to_s3


class SaveMarkdownToWordToolSchema(BaseModel):
    markdown_text: str = Field(..., description="The markdown-formatted text to convert.")
    filename: str = Field(..., description="The desired name for the output Word document.")
    img_directory: str = Field(..., description="The storage path where the document should be stored.")


def _render_document(markdown_text: str) -> Document:
    html = markdown(markdown_text, extensions=["tables", "fenced_code", "footnotes", "toc"])
    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    for element in soup.contents:
        name = getattr(element, "name", None)
        if not name:
            text = str(element).strip()
            if text:
                document.add_paragraph(text)
            continue
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = min(int(name[1]) - 1, 4)
            document.add_heading(element.get_text(strip=True), level=level)
            continue
        if name == "p":
            paragraph = document.add_paragraph()
            paragraph.add_run(element.get_text())
            paragraph.paragraph_format.space_after = Pt(12)
            continue
        if name in {"ul", "ol"}:
            style = "List Bullet" if name == "ul" else "List Number"
            for item in element.find_all("li", recursive=False):
                document.add_paragraph(item.get_text(" ", strip=True), style=style)
            continue
        if name == "blockquote":
            paragraph = document.add_paragraph()
            paragraph.style = "Intense Quote"
            paragraph.add_run(element.get_text(" ", strip=True))
            continue
        if name == "pre":
            paragraph = document.add_paragraph()
            paragraph.add_run(element.get_text())
            continue
        if name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            column_count = max(len(row.find_all(["th", "td"])) for row in rows)
            table = document.add_table(rows=0, cols=column_count)
            for row in rows:
                cells = row.find_all(["th", "td"])
                new_row = table.add_row().cells
                for index, cell in enumerate(cells):
                    new_row[index].text = cell.get_text(" ", strip=True)
            continue

        text = element.get_text(" ", strip=True)
        if text:
            document.add_paragraph(text)

    return document


def save_markdown_to_word(
        markdown_text: str,
        filename: str,
        img_directory: str,
        *,
        process_id: str | None = None,
        run_by: str | None = None,
        **_: Any,
) -> str:
    if not markdown_text:
        return "Error: No markdown text provided."
    if not filename:
        return "Error: No filename provided."

    try:
        document = _render_document(markdown_text)
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, filename)
            document.save(temp_file_path)
            result = upload_to_s3(
                img_directory,
                process_id,
                run_by,
                [temp_file_path],
                [filename],
            )
            if result.get("uploaded_files"):
                uploaded_path = result["uploaded_files"][0]
            else:
                uploaded_path = f"user_{run_by}/workflow_{img_directory}/run_{process_id}/{filename}"
            bucket_name = os.getenv("S3_BUCKET_NAME", "mybucket")
            return f"Document converted from markdown and uploaded to S3 at s3://{bucket_name}/{uploaded_path}."
    except Exception as exc:
        return f"Error in document conversion or upload process: {exc}"


__all__ = ["SaveMarkdownToWordToolSchema", "save_markdown_to_word"]
